import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = docx.Document()

# Page setup
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Custom Styles
def add_custom_style(name, type, font_name, size, color=None, bold=False):
    style = doc.styles.add_style(name, type)
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor(*color)
    font.bold = bold
    return style

# Add styles
add_custom_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Arial', 32, (0, 0, 0), True)
add_custom_style('SubtitleStyle', WD_STYLE_TYPE.PARAGRAPH, 'Arial', 18, (100, 100, 100))
add_custom_style('Heading1Style', WD_STYLE_TYPE.PARAGRAPH, 'Arial', 24, (0, 0, 0), True)
add_custom_style('Heading2Style', WD_STYLE_TYPE.PARAGRAPH, 'Arial', 18, (0, 200, 148), True) # Green accent
add_custom_style('BodyStyle', WD_STYLE_TYPE.PARAGRAPH, 'Calibri', 12, (50, 50, 50))
add_custom_style('QuoteStyle', WD_STYLE_TYPE.PARAGRAPH, 'Calibri', 14, (100, 100, 100), True)

# --- Slide 1: Cover ---
p = doc.add_paragraph(style='TitleStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SEO PLANET')

p = doc.add_paragraph('The New Era of Marketing', style='SubtitleStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('\nPremier Digital Marketing & SEO Agency\n\nEst. 2019 | 68+ Enterprise Clients', style='BodyStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# --- Slide 2: Philosophy ---
p = doc.add_paragraph('The SEO Planet Philosophy', style='Heading1Style')

p = doc.add_paragraph('We don\'t just chase algorithms; we build Enterprise Search Architecture & Growth Systems.\n', style='BodyStyle')
p = doc.add_paragraph('SEO Planet delivers enterprise-grade search architecture, high-performance acquisition, and advanced analytics to systematically accelerate your market presence and drive measurable growth.', style='BodyStyle')

p = doc.add_paragraph('\nOur Mission: To pair algorithmic SEO, performance ads, and content systems to help ambitious brands win.', style='QuoteStyle')

doc.add_page_break()

# --- Slide 3: Unified Growth Architecture ---
doc.add_paragraph('Our Unified Growth Architecture', style='Heading1Style')

doc.add_paragraph('Six core services. One unified growth architecture. Every engagement is built around measurable lift. We choose the right mix and you get a predictable trajectory.', style='BodyStyle')

doc.add_paragraph('\nCore Pillars', style='Heading2Style')
services = [
    "[S01] Enterprise Search",
    "[S02] Behavioral Analytics & CRO",
    "[S03] Generative Engine Optimization (GEO)",
    "[S04] Performance Marketing",
    "[S05] Topical Authority Content"
]
for s in services:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# --- Slide 4: Services Breakdown 1 ---
doc.add_paragraph('[S01] Enterprise Search', style='Heading1Style')
doc.add_paragraph('Advanced technical frameworks, entity-based search strategies, and comprehensive topical coverage.\n', style='BodyStyle')

p = doc.add_paragraph()
p.add_run('Our Flagship Offering:\n').bold = True
doc.add_paragraph('Technical SEO Mastery: Site architecture, log file analysis, core web vitals.', style='List Bullet')
doc.add_paragraph('Programmatic & Entity SEO: Schema markup, entity optimization.', style='List Bullet')
doc.add_paragraph('Proven Results: +340% Avg Traffic Lift, Top 3 SERP Targeting, AI-Ready Schema Stack.', style='List Bullet')

doc.add_paragraph('\n[S02] Behavioral Analytics & CRO', style='Heading1Style')
doc.add_paragraph('Systematic funnel teardowns, rigorous A/B testing, and quantitative research to optimize the user journey.\n', style='BodyStyle')
doc.add_paragraph('Traffic is only half the battle. We convert your visitors into revenue through Conversion Rate Optimization (CRO), A/B & Multivariate Testing, Funnel Analysis, and Behavioral Economics.', style='BodyStyle')

doc.add_page_break()

# --- Slide 5: Services Breakdown 2 ---
doc.add_paragraph('[S03] Generative Engine Optimization', style='Heading1Style')
doc.add_paragraph('Strategic adaptation ensuring brand visibility across LLMs and next-generation discovery platforms.\n', style='BodyStyle')
doc.add_paragraph('The future of search is Generative AI (ChatGPT, Perplexity, Gemini). Our services include LLM Readiness Assessment, Generative Engine Optimization, and NLP & Contextual Optimization.', style='BodyStyle')

doc.add_paragraph('\n[S04] & [S05] Growth Accelerators', style='Heading1Style')

doc.add_paragraph('Performance Marketing', style='Heading2Style')
doc.add_paragraph('Data-driven media buying optimized for predictable scaling and maximum efficiency across Google Ads, Meta Ads, LinkedIn Ads, and Programmatic Media Buying.', style='BodyStyle')

doc.add_paragraph('Topical Authority Content', style='Heading2Style')
doc.add_paragraph('Comprehensive content strategies that establish deep topical relevance and industry leadership via Content Engineering, Digital PR, and Brand Ecosystem Architecture.', style='BodyStyle')

doc.add_page_break()

# --- Slide 6: Contact ---
p = doc.add_paragraph('Let\'s Start Scaling', style='Heading1Style')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('\nReady to dominate your industry?', style='SubtitleStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('\nWe are currently booking for upcoming quarters. Partner with the agency that engineers predictable revenue growth.', style='QuoteStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('\n1. Audit: We perform a systematic teardown of your current architecture.\n2. Strategy: We design a custom unified growth architecture.\n3. Execution: We deploy our systems and scale your revenue.', style='BodyStyle')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('\nContact SEO Planet Today (seoplanet.in)', style='Heading2Style')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('c:\\Users\\yuppp\\Downloads\\Eventa-main\\SEO_Planet_Prospectus.docx')
print("Successfully generated SEO_Planet_Prospectus.docx")
